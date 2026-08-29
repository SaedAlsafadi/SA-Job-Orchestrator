import asyncio
from httpx import AsyncClient, ASGITransport
from app.main import app
from app.core.security import create_access_token

async def test():
    token = create_access_token("f9305c63796d4430bcdb178025ea6d64")
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        import docx
        from io import BytesIO
        doc = docx.Document()
        doc.add_paragraph("Saed Alsafadi")
        doc.add_paragraph("saed@email.com | +966 555 1234")
        doc.add_paragraph("")
        doc.add_paragraph("Software Engineer with 5 years of experience in Python, TypeScript, and cloud infrastructure.")
        doc.add_paragraph("")
        doc.add_paragraph("Experience")
        doc.add_paragraph("Senior Software Engineer at TechCorp (2021-Present)")
        doc.add_paragraph("- Led backend development using Python/FastAPI")
        doc.add_paragraph("- Designed microservices architecture on AWS")
        doc.add_paragraph("")
        doc.add_paragraph("Skills: Python, TypeScript, FastAPI, React, AWS, Docker")
        docx_bytes = BytesIO()
        doc.save(docx_bytes)

        files = {"file": ("resume.docx", docx_bytes.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        res = await client.post("/api/v1/candidate-profile/import-resume", files=files, headers={"Authorization": f"Bearer {token}"})
        print(f"STATUS: {res.status_code}")
        if res.status_code == 200:
            data = res.json()
            ident = data.get("identity", {})
            print(f"First Name: {ident.get('first_name', {})}")
            print(f"Last Name: {ident.get('last_name', {})}")
            print(f"Email: {ident.get('email', {})}")
            print(f"Summary: {str(ident.get('professional_summary', {}))[:100]}")
            print(f"Experience count: {len(data.get('experience', []))}")
            print(f"Skills count: {len(data.get('skills', []))}")
            if data.get("experience"):
                exp = data["experience"][0]
                print(f"  Exp[0] company: {exp.get('company', {})}")
                print(f"  Exp[0] title: {exp.get('title', {})}")
        else:
            print(res.text)

asyncio.run(test())
